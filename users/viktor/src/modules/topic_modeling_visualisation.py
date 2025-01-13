import os
import matplotlib.pyplot as plt
from wordcloud import WordCloud
import math

import subprocess
from interfaces import ITopicModelVisualizer
from datetime import datetime
from docx import Document

class TopicModelVisualizer(ITopicModelVisualizer):

    def __init__(self, model, output_subfolder, experiment_file):
        """
        Initializes the TopicModelVisualizer with the provided model, output folder, and experiment file.

        Args:
            model: The topic modeling algorithm or object to be visualized.
            output_subfolder (str): The folder where the generated visualizations and reports will be saved.
            experiment_file (str): The file containing the experiment details or configuration.
        """
        self.model = model
        self.output_subfolder = output_subfolder
        self.experiment_file = experiment_file
    
    def plot_inertia_graph(self, inertia_value):
        """
        Creates and saves a bar plot of the inertia value.

        The inertia value is visualized as a single bar indicating the sum of squared distances from 
        each point to its assigned cluster center.

        Args:
            inertia_value (float): The inertia value to be visualized.
        """
        plt.figure(figsize=(10, 6))
        plt.bar(["Inertia"], [inertia_value], color='blue')
        plt.title("Inertia of the Clustering Model", fontsize=14)
        plt.ylabel("Inertia", fontsize=12)
        plt.grid(axis='y', linestyle='--', alpha=0.7)
        graph_path = os.path.join(self.output_subfolder, "inertia_graph.png")
        plt.savefig(graph_path)
        plt.close()
        print(f"Inertia graph saved to: {graph_path}")

    def plot_ari_graph(self, ari_value):
        """
        Creates and saves a bar plot of the Adjusted Rand Index (ARI) value.

        The ARI value is visualized as a single bar indicating the similarity between the true and predicted clusters.

        Args:
            ari_value (float): The ARI value to be visualized.
        """
        plt.figure(figsize=(10, 6))
        plt.bar(["ARI"], [ari_value], color='orange')
        plt.title("Adjusted Rand Index (ARI)", fontsize=14)
        plt.ylabel("ARI", fontsize=12)
        plt.grid(axis='y', linestyle='--', alpha=0.7)
        graph_path = os.path.join(self.output_subfolder, "ari_graph.png")
        plt.savefig(graph_path)
        plt.close()
        print(f"ARI graph saved to: {graph_path}")

    def create_wordclouds(self, keywords):
        """
        Creates and saves word clouds for the generated topics.

        Args:
            keywords (dict): Dictionary of top keywords for each topic.
        """
        n_topics = len(keywords)
        topics_per_image = 20  # Number of topics per image
        n_images = math.ceil(n_topics / topics_per_image)

        for img_idx in range(n_images):
            # Determine the range of topics for this image
            start_idx = img_idx * topics_per_image
            end_idx = min(start_idx + topics_per_image, n_topics)
            topics_subset = list(keywords.items())[start_idx:end_idx]

            # Create a figure to display the word clouds
            n_cols = 5  # Number of columns for the grid layout
            n_rows = math.ceil(len(topics_subset) / n_cols)

            fig, axes = plt.subplots(n_rows, n_cols, figsize=(20, 4 * n_rows))
            axes = axes.flatten()  # Flatten axes for easy indexing

            for i, (topic, words) in enumerate(topics_subset):
                # Generate the word cloud
                wordcloud = WordCloud(width=800, height=400, background_color='white').generate(' '.join(words))

                # Display the word cloud in the subplot
                axes[i].imshow(wordcloud, interpolation='bilinear')
                axes[i].set_title(f"Topic {topic}", fontsize=14)
                axes[i].axis('off')

            # Turn off any remaining empty subplots
            for j in range(len(topics_subset), len(axes)):
                axes[j].axis('off')

            # Save the figure to the output folder
            wordcloud_file = os.path.join(self.output_subfolder, f"wordclouds_part_{img_idx + 1}.png")
            plt.tight_layout()
            plt.savefig(wordcloud_file, dpi=300)
            plt.close()

            print(f"Wordclouds saved to: {wordcloud_file}")

    def visualize_topic_diversity(self, topic_diversity):
        """
        Creates and saves a bar plot of the topic diversity score.

        The diversity score is visualized as a single bar showing the uniqueness of words 
        across all topics.

        Args:
            topic_diversity (float): The topic diversity score to be visualized.
        """
        plt.figure(figsize=(5, 5))
        plt.bar(['Diversity Score'], [topic_diversity], color='green')
        plt.ylabel('Diversity Score')
        plt.title('Topic Diversity')
        plt.savefig(os.path.join(self.output_subfolder, 'topic_diversity.png'))
        plt.close()
        print("Topic diversity visualization saved.")

    def visualize_silhouette_score(self, silhouette_score):
        """
        Creates and saves a bar plot of the silhouette score.

        The silhouette score is visualized as a single bar indicating the clustering quality.

        Args:
            silhouette_score (float): The silhouette score to be visualized.
        """
        plt.figure(figsize=(5, 5))
        plt.bar(['Silhouette Score'], [silhouette_score], color='orange')
        plt.ylabel('Silhouette Score')
        plt.title('Silhouette Score')
        plt.savefig(os.path.join(self.output_subfolder, 'silhouette_score.png'))
        plt.close()
        print("Silhouette score visualization saved.")

    def visualize_clustering_stability(self, clustering_stability):
        """
        Creates and saves a bar plot of the clustering stability score.

        The clustering stability score is visualized as a single bar indicating the stability 
        of the clustering results across multiple runs.

        Args:
            clustering_stability (float): The clustering stability score to be visualized.
        """
        plt.figure(figsize=(5, 5))
        plt.bar(['Clustering Stability'], [clustering_stability], color='purple')
        plt.ylabel('Stability Score')
        plt.title('Clustering Stability Score')
        plt.savefig(os.path.join(self.output_subfolder, 'clustering_stability.png'))
        plt.close()
        print("Clustering stability visualization saved.")

    def plot_topic_percentage(self, topic_percentages):
        """
        Visualizes the percentage of each topic in the dataset as a bar chart with different colors.

        Args:
            topic_percentages (dict): Dictionary with topics as keys and their corresponding percentages as values.
        """
        # Ensure the output folder exists
        os.makedirs(self.output_subfolder, exist_ok=True)
        
        # Extract topic names and their corresponding percentages
        topics = list(topic_percentages.keys())
        percentages = list(topic_percentages.values())
        
        # Set up the bar chart
        plt.figure(figsize=(10, 6))
        bars = plt.bar(topics, percentages, color=plt.cm.tab20.colors[:len(topics)])

        # Add labels and title
        plt.xlabel('Topics')
        plt.ylabel('Percentage (%)')
        plt.title('Topic Percentages in Dataset')
        
        # Rotate the x-axis labels by 90 degrees for readability
        plt.xticks(rotation=90)
        
        # Display percentage values on top of each bar
        for bar, percentage in zip(bars, percentages):
            yval = bar.get_height()
            plt.text(bar.get_x() + bar.get_width() / 2, yval, f"{percentage:.2f}%", ha='center', va='bottom')

        # Save and show the plot
        plt.tight_layout()
        output_path = os.path.join(self.output_subfolder, 'topic_percentage_chart.png')
        plt.savefig(output_path)
        
        print(f"Topic percentage chart saved to {output_path}")

    def get_git_info(self):
        """
        Get Git repository information using subprocess (for version and repository name).

        Returns:
            tuple: Git version and repository name if available, else None, None.
        """
        try:
            # Get Git version
            git_version = subprocess.check_output(['git', 'describe', '--always']).decode().strip()
            # Get Git repository name (assuming the repository is a git repository)
            git_repo_name = subprocess.check_output(['git', 'config', '--get', 'remote.origin.url']).decode().strip()
            git_repo_name = os.path.basename(git_repo_name).replace('.git', '')  # Extract repo name
            return git_version, git_repo_name
        except subprocess.CalledProcessError:
            return None, None  # In case Git is not available or fails

    def generate_report(self, topic_diversity, silhouette_score, clustering_stability, topic_percentages):
        """
        Generates a Word document report with all the experiment details, metrics, and Git information.

        Args:
            topic_diversity (float): The topic diversity score.
            silhouette_score (float): The silhouette score.
            clustering_stability (float): The clustering stability score.
            topic_percentages (dict): Dictionary with topics as keys and their corresponding percentages as values.
        """
        # Create the Word document
        doc = Document()

        # Add title and introduction
        doc.add_heading('Topic Model Experiment Report', 0)
        doc.add_paragraph('This report summarizes the results of the Topic Modeling experiment.')

        # Add experiment details
        doc.add_heading('Experiment Details', level=1)
        doc.add_paragraph(f'Experiment File: {self.experiment_file}')
        doc.add_paragraph(f'Execution Time: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        
        # Git Information
        git_version, git_repo_name = self.get_git_info()
        if git_version and git_repo_name:
            doc.add_heading('Git Information', level=1)
            doc.add_paragraph(f'Git Repository: {git_repo_name}')
            doc.add_paragraph(f'Git Version: {git_version}')
        else:
            doc.add_paragraph('Git information not available (Git is not installed or not initialized in the current directory).')

        # Add metrics details
        doc.add_heading('Metrics', level=1)
        doc.add_paragraph(f'Topic Diversity Score: {topic_diversity}')
        doc.add_paragraph(f'Silhouette Score: {silhouette_score}')
        doc.add_paragraph(f'Clustering Stability Score: {clustering_stability}')

        # Add topic percentages
        doc.add_heading('Topic Percentages', level=1)
        for topic, percentage in topic_percentages.items():
            doc.add_paragraph(f'Topic {topic}: {percentage:.2f}%')

        # Add the general explanation of the experiment
        doc.add_heading('Experiment Process', level=1)
        doc.add_paragraph(
            'This experiment involves applying topic modeling to a dataset and '
            'visualizing various metrics related to the topic modeling. The experiment also includes the calculation '
            'of topic diversity, silhouette score, clustering stability, and cosine similarity matrix.' 
        )

        # Save the report
        report_path = os.path.join(self.output_subfolder, 'topic_modeling_experiment_report.docx')
        doc.save(report_path)
        print(f'Report saved to {report_path}')

    def create_all_visualizations(self, topic_diversity, silhouette_score, clustering_stability, topic_percentages, inertia, ari_score, keywords):
        """
        Generates and saves all visualizations and then creates a report with the results.

        Args:
            topic_diversity (float): The topic diversity score.
            silhouette_score (float): The silhouette score.
            clustering_stability (float): The clustering stability score.
            topic_percentages (dict): Dictionary with topics as keys and their corresponding percentages as values.
            inertia (float): The inertia value of the clustering model.
            ari_score (float): The Adjusted Rand Index (ARI) score.
            keywords (dict): Dictionary of top keywords for each topic.
        """
        self.visualize_silhouette_score(silhouette_score)
        self.visualize_clustering_stability(clustering_stability)
        self.plot_topic_percentage(topic_percentages)
        self.create_wordclouds(keywords)
        self.visualize_topic_diversity(topic_diversity)
        self.plot_inertia_graph(inertia)
        self.plot_ari_graph(ari_score)

        # Generate the final report
        self.generate_report(topic_diversity, silhouette_score, clustering_stability, topic_percentages)