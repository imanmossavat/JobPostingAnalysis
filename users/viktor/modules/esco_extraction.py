import pandas as pd
import datetime
import os
import subprocess
from docx import Document
from interfaces.skill_knowledge_extractor import ISkillKnowledgeExtractor
from transformers import pipeline
from langdetect import detect, LangDetectException


class ESCOAnalyzer(ISkillKnowledgeExtractor):
    def __init__(self):
        # Initialize the skill and knowledge extraction pipelines
        self.token_skill_classifier = pipeline(
            model="jjzha/escoxlmr_skill_extraction",
            aggregation_strategy="first"
        )
        self.token_knowledge_classifier = pipeline(
            model="jjzha/escoxlmr_knowledge_extraction",
            aggregation_strategy="first"
        )

    def extract_skills_and_knowledge(self, text: str, lang: str) -> dict:
        """
        Extracts skills and knowledge entities from the given text.

        Args:
            text (str): The input text to analyze.
            lang (str): The language of the input text.

        Returns:
            dict: A dictionary containing the text, detected skills, knowledge, and detected language.
        """
        output_skills = self.token_skill_classifier(text)
        output_knowledge = self.token_knowledge_classifier(text)

        # Process skills
        for result in output_skills:
            if result.get("entity_group"):
                result["entity"] = "Skill"
                del result["entity_group"]

        # Process knowledge
        for result in output_knowledge:
            if result.get("entity_group"):
                result["entity"] = "Knowledge"
                del result["entity_group"]

        # Aggregate spans (tokens that should be merged together)
        output_skills = self.aggregate_span(output_skills, text)
        output_knowledge = self.aggregate_span(output_knowledge, text)

        return {
            "text": text,
            "skills": output_skills,
            "knowledge": output_knowledge,
            "detected-language": lang
        }

    def save_progress(self, progress_data: pd.DataFrame, output_subfolder: str) -> None:
        """
        Saves the progress of the extraction to a CSV file.

        Args:
            progress_data (pd.DataFrame): Data to be saved.
            output_subfolder (str): Folder where to save the progress file.
        """
        timestamp = datetime.datetime.now().strftime('%Y%m%d_%H%M%S')
        progress_file = os.path.join(output_subfolder, f'extracted_skills_{timestamp}.csv')
        progress_data.to_csv(progress_file, index=False, encoding='utf-8-sig')
        print(f"Progress saved to {progress_file}")

    def generate_report(self, input_file: str, output_file: str, output_subfolder: str) -> None:
        """
        Generates a Word document report summarizing the experiment.

        Args:
            input_file (str): The input file path.
            output_file (str): The output file path.
            output_subfolder (str): Folder where the report will be saved.
        """
        try:
            git_version = subprocess.check_output(["git", "rev-parse", "HEAD"], cwd=os.path.dirname(__file__)).strip().decode()
        except Exception:
            git_version = "Git version not available"

        timestamp = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        os.makedirs(output_subfolder, exist_ok=True)
        report_file = os.path.join(output_subfolder, "experiment_report.docx")

        # Create the Word document
        doc = Document()
        doc.add_heading("Experiment Report", level=1)
        doc.add_paragraph(f"Timestamp: {timestamp}")
        doc.add_paragraph(f"Input File: {input_file}")
        doc.add_paragraph(f"Output File: {output_file}")
        doc.add_paragraph(f"Output Subfolder: {output_subfolder}")
        doc.add_paragraph(f"Git Version: {git_version}")

        doc.add_heading("Pipeline Explanation", level=2)
        doc.add_paragraph(
            "This experiment extracts skills and knowledge from job descriptions using pre-trained NER models."
        )

        doc.add_heading("Progress and Final Results", level=2)
        doc.add_paragraph(
            f"Final results are saved to {output_file}. The extraction process was completed successfully."
        )

        doc.save(report_file)
        print(f"Experiment report saved to: {report_file}")

    def aggregate_span(self, results, text):
        """
        Aggregates consecutive token spans into single entities.
        """
        # Your implementation for aggregating spans goes here
        return results


def detect_language(text: str) -> str:
    """
    Detects the language of the input text.
    """
    try:
        lang = detect(text)
        return 'english' if lang == 'en' else 'dutch' if lang == 'nl' else 'Other'
    except LangDetectException:
        return 'unknown'


def initiate_esco_analysis(input_file: str, output_file: str, output_subfolder: str):
    """
    Runs the ESCO analysis for extracting skills and knowledge from job descriptions.
    """
    analyzer = ESCOAnalyzer()
    df = pd.read_csv(input_file, encoding='utf-8')
    results = []

    for i, row in df.iterrows():
        job_description = row["Description"]
        lang = detect_language(job_description)
        extracted_data = analyzer.extract_skills_and_knowledge(job_description, lang)
        results.append(extracted_data)

        if (i + 1) % (len(df) // 10) == 0 or i == len(df) - 1:
            progress_df = pd.DataFrame(results)
            analyzer.save_progress(progress_df, output_subfolder)

    # Final save
    final_df = pd.DataFrame(results)
    final_df.to_csv(output_file, index=False, encoding='utf-8-sig')
    analyzer.generate_report(input_file, output_file, output_subfolder)